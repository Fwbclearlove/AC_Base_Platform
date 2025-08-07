import os
import pandas as pd
import pdfplumber
import PyPDF2
import docx
from docx import Document
import fitz  # PyMuPDF
from tabulate import tabulate
import time
from typing import List, Dict, Any


class DataExtractor:
    def __init__(self, pdf_path: str, docx_path: str):
        self.pdf_path = pdf_path
        self.docx_path = docx_path
        self.results = {
            'pdf': {'tables': [], 'text': '', 'extraction_time': 0, 'error': None},
            'docx': {'tables': [], 'text': '', 'extraction_time': 0, 'error': None}
        }

    def extract_from_pdf(self):
        """使用多种方法从PDF提取数据"""
        start_time = time.time()

        try:
            # 方法1: 使用pdfplumber提取表格（推荐）
            print("📄 正在使用 pdfplumber 提取PDF数据...")
            tables = []
            full_text = ""

            with pdfplumber.open(self.pdf_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    print(f"   处理第 {i + 1} 页...")

                    # 提取文本
                    page_text = page.extract_text()
                    if page_text:
                        full_text += f"\n--- 第{i + 1}页 ---\n{page_text}\n"

                    # 提取表格
                    page_tables = page.extract_tables()
                    for j, table in enumerate(page_tables):
                        if table and len(table) > 1:
                            try:
                                # 清理表格数据
                                cleaned_table = []
                                for row in table:
                                    cleaned_row = [cell.strip() if cell else '' for cell in row]
                                    cleaned_table.append(cleaned_row)

                                df = pd.DataFrame(cleaned_table[1:], columns=cleaned_table[0])
                                tables.append({
                                    'page': i + 1,
                                    'table_index': j + 1,
                                    'data': df,
                                    'shape': df.shape,
                                    'source': 'pdfplumber'
                                })
                            except Exception as e:
                                print(f"      表格 {j + 1} 处理失败: {str(e)}")

            # 备用方法: PyMuPDF (如果pdfplumber效果不好)
            print("📄 使用 PyMuPDF 作为备用方法...")
            doc = fitz.open(self.pdf_path)
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                # 查找表格
                tabs = page.find_tables()
                for tab_index, tab in enumerate(tabs):
                    try:
                        df = tab.to_pandas()
                        if not df.empty:
                            tables.append({
                                'page': page_num + 1,
                                'table_index': tab_index + 1,
                                'data': df,
                                'shape': df.shape,
                                'source': 'PyMuPDF'
                            })
                    except Exception as e:
                        print(f"      PyMuPDF表格 {tab_index + 1} 处理失败: {str(e)}")

            self.results['pdf']['tables'] = tables
            self.results['pdf']['text'] = full_text
            self.results['pdf']['extraction_time'] = time.time() - start_time

        except Exception as e:
            self.results['pdf']['error'] = str(e)
            self.results['pdf']['extraction_time'] = time.time() - start_time
            print(f"❌ PDF提取失败: {str(e)}")

    def extract_from_docx(self):
        """从DOCX文件提取数据"""
        start_time = time.time()

        try:
            print("📝 正在从DOCX提取数据...")
            doc = Document(self.docx_path)

            # 提取文本
            full_text = ""
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    full_text += paragraph.text + "\n"

            # 提取表格
            tables = []
            for i, table in enumerate(doc.tables):
                try:
                    # 提取表格数据
                    data = []
                    for row in table.rows:
                        row_data = []
                        for cell in row.cells:
                            # 清理单元格文本
                            cell_text = cell.text.strip().replace('\n', ' ')
                            row_data.append(cell_text)
                        data.append(row_data)

                    if len(data) > 1 and any(any(cell for cell in row) for row in data):
                        # 检查是否有有效的表头
                        headers = data[0] if data[0] and any(h.strip() for h in data[0]) else [f"列{j + 1}" for j in
                                                                                               range(len(data[0]))]

                        df = pd.DataFrame(data[1:], columns=headers)
                        # 移除完全空白的行和列
                        df = df.dropna(how='all').dropna(axis=1, how='all')

                        if not df.empty:
                            tables.append({
                                'page': 'N/A',
                                'table_index': i + 1,
                                'data': df,
                                'shape': df.shape,
                                'source': 'python-docx'
                            })

                except Exception as e:
                    print(f"   表格 {i + 1} 处理失败: {str(e)}")

            self.results['docx']['tables'] = tables
            self.results['docx']['text'] = full_text
            self.results['docx']['extraction_time'] = time.time() - start_time

        except Exception as e:
            self.results['docx']['error'] = str(e)
            self.results['docx']['extraction_time'] = time.time() - start_time
            print(f"❌ DOCX提取失败: {str(e)}")

    def compare_results(self):
        """对比提取结果"""
        print("\n" + "=" * 80)
        print("📊 提取结果对比分析")
        print("=" * 80)

        # 基本统计对比
        comparison_data = []

        for format_type in ['pdf', 'docx']:
            result = self.results[format_type]

            if result['error']:
                comparison_data.append([
                    format_type.upper(),
                    "❌ 提取失败",
                    f"错误: {result['error'][:50]}...",
                    f"{result['extraction_time']:.2f}s"
                ])
            else:
                text_length = len(result['text'])
                table_count = len(result['tables'])
                total_cells = sum(table['data'].size for table in result['tables'])

                comparison_data.append([
                    format_type.upper(),
                    f"✅ 成功",
                    f"表格数: {table_count}, 文本长度: {text_length}, 总单元格: {total_cells}",
                    f"{result['extraction_time']:.2f}s"
                ])

        print(tabulate(comparison_data,
                       headers=['格式', '状态', '提取内容', '耗时'],
                       tablefmt='grid'))

        # 详细表格对比
        self._compare_tables()

        # 文本对比
        self._compare_text()

        # 给出建议
        self._give_recommendations()

    def _compare_tables(self):
        """对比表格提取效果"""
        print(f"\n📋 表格提取详细对比:")
        print("-" * 60)

        pdf_tables = self.results['pdf'].get('tables', [])
        docx_tables = self.results['docx'].get('tables', [])

        print(f"PDF表格数量: {len(pdf_tables)}")
        print(f"DOCX表格数量: {len(docx_tables)}")

        # 显示PDF表格详情
        if pdf_tables:
            print(f"\n📄 PDF表格详情:")
            for i, table in enumerate(pdf_tables):
                print(
                    f"  表格{i + 1} (第{table['page']}页, {table['source']}): {table['shape'][0]}行 x {table['shape'][1]}列")
                if table['shape'][0] <= 5:  # 小表格显示内容
                    print("    预览:")
                    print(table['data'].to_string(index=False).replace('\n', '\n    '))
                print()

        # 显示DOCX表格详情
        if docx_tables:
            print(f"📝 DOCX表格详情:")
            for i, table in enumerate(docx_tables):
                print(f"  表格{i + 1}: {table['shape'][0]}行 x {table['shape'][1]}列")
                if table['shape'][0] <= 5:  # 小表格显示内容
                    print("    预览:")
                    print(table['data'].to_string(index=False).replace('\n', '\n    '))
                print()

    def _compare_text(self):
        """对比文本提取效果"""
        print(f"📝 文本提取对比:")
        print("-" * 60)

        pdf_text_len = len(self.results['pdf'].get('text', ''))
        docx_text_len = len(self.results['docx'].get('text', ''))

        print(f"PDF文本长度: {pdf_text_len} 字符")
        print(f"DOCX文本长度: {docx_text_len} 字符")

        # 显示文本开头预览
        if pdf_text_len > 0:
            preview = self.results['pdf']['text'][:200]
            print(f"\nPDF文本预览:\n{preview}...")

        if docx_text_len > 0:
            preview = self.results['docx']['text'][:200]
            print(f"\nDOCX文本预览:\n{preview}...")

    def _give_recommendations(self):
        """给出使用建议"""
        print(f"\n💡 使用建议:")
        print("-" * 60)

        pdf_tables = len(self.results['pdf'].get('tables', []))
        docx_tables = len(self.results['docx'].get('tables', []))
        pdf_error = self.results['pdf'].get('error')
        docx_error = self.results['docx'].get('error')

        if pdf_error and docx_error:
            print("❌ 两种格式都提取失败，建议检查文件完整性")
        elif pdf_error:
            print("✅ 建议使用DOCX格式，PDF提取失败")
        elif docx_error:
            print("✅ 建议使用PDF格式，DOCX提取失败")
        elif docx_tables > pdf_tables:
            print("✅ 建议使用DOCX格式，表格提取数量更多")
        elif pdf_tables > docx_tables:
            print("✅ 建议使用PDF格式，表格提取数量更多")
        else:
            pdf_time = self.results['pdf']['extraction_time']
            docx_time = self.results['docx']['extraction_time']
            if pdf_time < docx_time:
                print("✅ 两种格式效果相近，PDF处理速度更快")
            else:
                print("✅ 两种格式效果相近，DOCX处理速度更快")

    def save_results(self, output_dir="extraction_results"):
        """保存提取结果"""
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        # 保存表格
        for format_type in ['pdf', 'docx']:
            tables = self.results[format_type].get('tables', [])
            for i, table in enumerate(tables):
                filename = f"{output_dir}/{format_type}_table_{i + 1}.csv"
                table['data'].to_csv(filename, index=False, encoding='utf-8-sig')
                print(f"💾 已保存: {filename}")

        # 保存文本
        for format_type in ['pdf', 'docx']:
            text = self.results[format_type].get('text', '')
            if text:
                filename = f"{output_dir}/{format_type}_text.txt"
                with open(filename, 'w', encoding='utf-8') as f:
                    f.write(text)
                print(f"💾 已保存: {filename}")


def main():
    # 配置文件路径
    PDF_PATH =  r"C:\Users\86152\Desktop\RuifanLi 老师论文\多视图有监督的LDA模型.pdf" # 修改为你的PDF文件路径
    DOCX_PATH = r"C:\Users\86152\Desktop\学生论文\多视图有监督的LDA模型.docx"  # 修改为你的DOCX文件路径

    # 检查文件是否存在
    if not os.path.exists(PDF_PATH):
        print(f"❌ 找不到PDF文件: {PDF_PATH}")
        return

    if not os.path.exists(DOCX_PATH):
        print(f"❌ 找不到DOCX文件: {DOCX_PATH}")
        return

    print("🚀 开始数据提取效果对比...")
    print(f"📄 PDF文件: {PDF_PATH}")
    print(f"📝 DOCX文件: {DOCX_PATH}")
    print("-" * 80)

    # 创建提取器
    extractor = DataExtractor(PDF_PATH, DOCX_PATH)

    # 执行提取
    extractor.extract_from_pdf()
    extractor.extract_from_docx()

    # 对比结果
    extractor.compare_results()

    # 保存结果
    save_results = input("\n💾 是否保存提取结果到文件? (y/n): ").lower() == 'y'
    if save_results:
        extractor.save_results()

    print("\n✅ 对比完成！")


if __name__ == "__main__":
    # 检查依赖库
    required_libs = ['pdfplumber', 'PyPDF2', 'python-docx', 'PyMuPDF', 'tabulate', 'pandas']
    missing_libs = []

    for lib in required_libs:
        try:
            if lib == 'python-docx':
                import docx
            elif lib == 'PyMuPDF':
                import fitz
            else:
                __import__(lib)
        except ImportError:
            missing_libs.append(lib)

    if missing_libs:
        print("❌ 缺少以下依赖库，请先安装:")
        for lib in missing_libs:
            print(f"   pip install {lib}")
        print()
    else:
        main()